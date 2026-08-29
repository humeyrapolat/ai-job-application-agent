import base64
import binascii
from io import BytesIO
from pathlib import Path

MAX_DOCUMENT_BYTES = 5 * 1024 * 1024
SUPPORTED_EXTENSIONS = {".docx", ".md", ".pdf", ".txt"}


class DocumentTextExtractionError(Exception):
    """Raised when uploaded document text cannot be extracted."""


def extract_document_text(*, filename: str, content_base64: str) -> str:
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise DocumentTextExtractionError(
            "Supported CV file types are PDF, DOCX, TXT, and MD."
        )

    content = _decode_base64_document(content_base64)
    if len(content) > MAX_DOCUMENT_BYTES:
        raise DocumentTextExtractionError("CV file is too large. Please upload a file under 5 MB.")

    if extension in {".txt", ".md"}:
        text = _decode_plain_text(content)
    elif extension == ".pdf":
        text = _extract_pdf_text(content)
    else:
        text = _extract_docx_text(content)

    cleaned_text = _clean_text(text)
    if len(cleaned_text) < 20:
        raise DocumentTextExtractionError(
            "The CV file did not contain enough readable text to analyze."
        )

    return cleaned_text


def _decode_base64_document(content_base64: str) -> bytes:
    try:
        return base64.b64decode(content_base64, validate=True)
    except (binascii.Error, ValueError) as exc:
        raise DocumentTextExtractionError("CV file content could not be decoded.") from exc


def _decode_plain_text(content: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue

    raise DocumentTextExtractionError("Text CV encoding is not supported.")


def _extract_pdf_text(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentTextExtractionError("PDF CV support is not installed.") from exc

    try:
        reader = PdfReader(BytesIO(content))
        if reader.is_encrypted:
            reader.decrypt("")
        page_text = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:
        raise DocumentTextExtractionError("PDF CV text could not be extracted.") from exc

    return "\n".join(page_text)


def _extract_docx_text(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentTextExtractionError("DOCX CV support is not installed.") from exc

    try:
        document = Document(BytesIO(content))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
    except Exception as exc:
        raise DocumentTextExtractionError("DOCX CV text could not be extracted.") from exc

    return "\n".join(parts)


def _clean_text(text: str) -> str:
    lines = (" ".join(line.split()) for line in text.splitlines())
    return "\n".join(line for line in lines if line).strip()
